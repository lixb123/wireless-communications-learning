from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\30534\Desktop\通信学习\业余无线电\地面站")
SRC = ROOT / "业余无线电卫星UV通信地面站建设方案及功能介绍_图示增强版.docx"
OUT = ROOT / "UV地面站方案.docx"
ASSET = ROOT / "方案图示素材"
BLUE="2E74B5"; NAVY="163A5F"; DARK="243444"; MUTED="667788"


def set_run(run,size=10.5,bold=False,color=DARK,italic=False):
    run.font.name="Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic
    run.font.color.rgb=RGBColor.from_string(color)


def set_alt(shape,text):
    shape._inline.docPr.set("descr",text); shape._inline.docPr.set("title",text)


def insert_after(ref, element):
    ref._p.addnext(element)


def new_picture_paragraph(doc,path,width,alt,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(2)
    shape=p.add_run().add_picture(str(path),width=Inches(width)); set_alt(shape,alt)
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before=Pt(2); cap.paragraph_format.space_after=Pt(7)
    set_run(cap.add_run(caption),8.5,False,MUTED,True)
    p_el=deepcopy(p._p); cap_el=deepcopy(cap._p)
    p._element.getparent().remove(p._element); cap._element.getparent().remove(cap._element)
    return p_el,cap_el


def add_picture_after(doc,ref,path,width,alt,caption):
    p,cap=new_picture_paragraph(doc,path,width,alt,caption)
    ref._p.addnext(cap); ref._p.addnext(p)


def set_cell_margins(cell,top=90,start=110,bottom=90,end=110):
    tcpr=cell._tc.get_or_add_tcPr(); mar=tcpr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); tcpr.append(mar)
    for tag,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        e=mar.find(qn(f"w:{tag}"))
        if e is None: e=OxmlElement(f"w:{tag}"); mar.append(e)
        e.set(qn("w:w"),str(val)); e.set(qn("w:type"),"dxa")


def set_borders(table,color="CBD6DF"):
    pr=table._tbl.tblPr; bd=pr.find(qn("w:tblBorders"))
    if bd is None: bd=OxmlElement("w:tblBorders"); pr.append(bd)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e=OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"6"); e.set(qn("w:color"),color); bd.append(e)


def gallery_table(doc,items):
    rows=(len(items)+1)//2
    t=doc.add_table(rows=rows,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; set_borders(t)
    for idx,(title,desc,filename) in enumerate(items):
        c=t.cell(idx//2,idx%2); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; set_cell_margins(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(1)
        sh=p.add_run().add_picture(str(ASSET/filename),width=Inches(2.25)); set_alt(sh,title+"典型外观示意")
        hp=c.add_paragraph(); hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; hp.paragraph_format.space_after=Pt(1)
        set_run(hp.add_run(title),9.4,True,NAVY)
        dp=c.add_paragraph(); dp.paragraph_format.space_after=Pt(0); dp.paragraph_format.line_spacing=1.0
        set_run(dp.add_run(desc),7.7)
    return t


def new_gallery_element(doc,items):
    t=gallery_table(doc,items)
    el=deepcopy(t._tbl); t._element.getparent().remove(t._element)
    return el


def paragraph_by_prefix(doc,prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix): return p
    raise ValueError(prefix)


def first_following_body(doc,heading):
    found=False
    for p in doc.paragraphs:
        if p is heading: found=True; continue
        if found and p.text.strip(): return p
    return heading


RF_ITEMS=[
    ("合路器/分路器","按频段合并或分离射频通道。","03_合路器分路器.png"),
    ("低噪声放大器（LNA）","提升弱信号，关注噪声、增益和过载。","04_低噪声放大器LNA.png"),
    ("功率放大器（PA）","受控放大发射激励并配合散热与联锁。","05_功率放大器PA.png"),
    ("射频滤波器","抑制带外干扰、谐波和共址信号。","06_射频滤波器.png"),
    ("射频同轴线缆与接头","按频段、长度和损耗统一核算。","07_射频同轴线缆与接头.png"),
    ("馈线浪涌保护器","馈线入室处泄放浪涌并可靠接地。","12_馈线浪涌保护器.png"),
]

CONTROL_ITEMS=[
    ("双频收发电台","完成调谐、调制解调和受控收发。","08_双频收发电台.png"),
    ("业务控制电脑","运行轨道、跟踪、多普勒和数据软件。","09_业务控制电脑.png"),
    ("旋转器控制器","执行角度指令并返回位置与限位。","10_旋转器控制器.png"),
    ("控制与供电线缆","承载驱动、反馈、控制和供电信号。","11_控制与供电线缆.png"),
]


def build():
    doc=Document(SRC)

    # Remove the former visual appendices C-E; their content will be redistributed into the body.
    appendix_c=paragraph_by_prefix(doc,"附录 C  系统示意图与模块实物图册")
    body=doc._element.body
    start=list(body).index(appendix_c._p)
    for el in list(body)[start:]:
        if el.tag != qn("w:sectPr"):
            body.remove(el)

    # 2.1: retain existing full-site photo; add a focused rotor photograph after its first descriptive paragraph.
    h21=paragraph_by_prefix(doc,"2.1 现场可见构成")
    body21=first_following_body(doc,h21)
    add_picture_after(doc,body21,ASSET/"02_方位俯仰旋转器_现场实景.png",4.55,
                      "方位俯仰旋转器现场局部实景","图 2-2  方位/俯仰旋转器及馈线局部实景")

    # 2.2: overall architecture diagram directly under the heading.
    h22=paragraph_by_prefix(doc,"2.2 总体架构")
    add_picture_after(doc,h22,ASSET/"系统示意图_01_总体架构.png",6.35,
                      "地面站总体系统架构示意图","图 2-3  地面站总体系统架构示意图")

    # 2.3: RF chain diagram before the process text.
    h23=paragraph_by_prefix(doc,"2.3 信号与控制关系")
    add_picture_after(doc,h23,ASSET/"系统示意图_02_射频链路.png",6.35,
                      "UV地面站射频接收和受控发射链路示意图","图 2-4  射频接收与受控发射链路示意图")

    # Section 3: two compact, page-break-controlled equipment plates.
    h3=paragraph_by_prefix(doc,"3  系统组成及接口说明")
    intro=doc.add_paragraph(); intro.paragraph_format.space_after=Pt(5)
    set_run(intro.add_run("以下图片用于说明各功能模块的典型实物形态；除现场照片外，不代表具体品牌、型号、尺寸、接口数量或最终供货形态。"),9.0,False,MUTED,True)
    intro_el=deepcopy(intro._p); intro._element.getparent().remove(intro._element)
    rf_title=doc.add_paragraph(style="Heading 2"); set_run(rf_title.add_run("3.1 射频链路模块典型外观"),13,True,BLUE); rf_title.paragraph_format.page_break_before=True
    rf_title_el=deepcopy(rf_title._p); rf_title._element.getparent().remove(rf_title._element)
    rf_el=new_gallery_element(doc,RF_ITEMS)
    ctrl_title=doc.add_paragraph(style="Heading 2"); set_run(ctrl_title.add_run("3.2 收发与控制模块典型外观"),13,True,BLUE); ctrl_title.paragraph_format.page_break_before=True
    ctrl_title_el=deepcopy(ctrl_title._p); ctrl_title._element.getparent().remove(ctrl_title._element)
    ctrl_el=new_gallery_element(doc,CONTROL_ITEMS)
    # Insert in reverse so final order is intro -> RF title/table -> control title/table -> original section-3 content.
    h3._p.addnext(ctrl_el); h3._p.addnext(ctrl_title_el); h3._p.addnext(rf_el); h3._p.addnext(rf_title_el); h3._p.addnext(intro_el)

    # Existing subsection must follow new visual subsections numerically.
    try:
        old=paragraph_by_prefix(doc,"3.1 接口设计要求")
        for run in old.runs: run.text=run.text.replace("3.1 接口设计要求","3.3 接口设计要求")
    except ValueError: pass

    # 4.1: tracking diagram next to its functional explanation.
    h41=paragraph_by_prefix(doc,"4.1 自动跟踪功能说明")
    body41=first_following_body(doc,h41)
    add_picture_after(doc,body41,ASSET/"系统示意图_03_跟踪控制.png",6.35,
                      "自动跟踪与设备控制关系示意图","图 4-1  自动跟踪与设备控制关系示意图")

    doc.core_properties.title="UV地面站方案"
    doc.save(OUT)
    print(OUT)


if __name__=="__main__": build()

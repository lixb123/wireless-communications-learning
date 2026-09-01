from pathlib import Path
from math import cos, sin, pi
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageOps, ImageEnhance
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\30534\Desktop\通信学习\业余无线电\地面站")
SRC = ROOT / "业余无线电卫星UV通信地面站建设方案及功能介绍.docx"
OUT = ROOT / "业余无线电卫星UV通信地面站建设方案及功能介绍_图示增强版.docx"
ASSET = ROOT / "方案图示素材"
ASSET.mkdir(exist_ok=True)
PHOTO1 = Path(r"F:\Wechat\xwechat_files\wxid_4k7n1aru2e4722_4016\temp\RWTemp\2026-08\19ab911195001df06d3ee996a8932cd2\1c7b7333ba7d4a62ffaa5c348835f841.jpg")
PHOTO2 = Path(r"F:\Wechat\xwechat_files\wxid_4k7n1aru2e4722_4016\temp\RWTemp\2026-08\19ab911195001df06d3ee996a8932cd2\9ed28b71e25901d842a957f869b20d63.jpg")
FONT = r"C:\Windows\Fonts\msyh.ttc"
FONT_B = r"C:\Windows\Fonts\msyhbd.ttc"

NAVY = "163A5F"; BLUE = "2E74B5"; DARK = "243444"; MUTED = "667788"
PALE = "EAF2F8"; LIGHT = "F4F6F9"; GREEN = "2B6E5B"; GOLD = "936F18"

def fnt(sz, bold=False):
    return ImageFont.truetype(FONT_B if bold else FONT, sz)

def rgb(h): return tuple(int(h[i:i+2],16) for i in (0,2,4))

def rounded(draw, box, fill, outline=None, width=2, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)

def shadow_card(w=900, h=600):
    bg = Image.new("RGB", (w,h), "#EDF2F5")
    sh = Image.new("RGBA", (w,h), (0,0,0,0))
    sd = ImageDraw.Draw(sh)
    sd.rounded_rectangle((82,62,w-62,h-52), radius=34, fill=(18,40,60,52))
    sh = sh.filter(ImageFilter.GaussianBlur(18))
    bg.paste(sh, (0,0), sh)
    d = ImageDraw.Draw(bg)
    rounded(d, (62,42,w-82,h-72), "#FAFCFD", "#CAD6E0", 3, 30)
    return bg, d

def badge(d, text="典型外观示意"):
    bbox=d.textbbox((0,0),text,font=fnt(23,True)); tw=bbox[2]-bbox[0]
    rounded(d,(82,62,82+tw+34,105),"#EAF2F8",None,1,18)
    d.text((99,70),text,font=fnt(23,True),fill="#2E74B5")

def save_card(img, name):
    path=ASSET/name
    img.save(path, quality=94)
    return path

def photo_card(src, crop, name, badge_text="现场实景提取"):
    bg,d=shadow_card(); badge(d,badge_text)
    im=Image.open(src).convert("RGB")
    x1,y1,x2,y2=crop
    im=im.crop((int(im.width*x1),int(im.height*y1),int(im.width*x2),int(im.height*y2)))
    im=ImageOps.autocontrast(im,cutoff=1)
    target=(720,405)
    im=ImageOps.fit(im,target,method=Image.Resampling.LANCZOS)
    mask=Image.new("L",target,0); md=ImageDraw.Draw(mask); md.rounded_rectangle((0,0,*target),radius=22,fill=255)
    bg.paste(im,(90,125),mask)
    return save_card(bg,name)

def connector(d,x,y,r=18,color="#D9C9A5"):
    d.ellipse((x-r,y-r,x+r,y+r),fill=color,outline="#6F6F6F",width=3)
    d.ellipse((x-r//2,y-r//2,x+r//2,y+r//2),fill="#EEF1F3",outline="#777",width=2)

def box3d(d, box, face="#B8C4CC", side="#8D9AA3", top="#DDE4E8", depth=30):
    x1,y1,x2,y2=box
    d.polygon([(x1,y1),(x1+depth,y1-depth),(x2+depth,y1-depth),(x2,y1)],fill=top,outline="#586772")
    d.polygon([(x2,y1),(x2+depth,y1-depth),(x2+depth,y2-depth),(x2,y2)],fill=side,outline="#586772")
    d.rounded_rectangle(box,radius=14,fill=face,outline="#586772",width=3)

def combiner_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(205,190,690,425),"#BFC9D0","#8E9AA3","#E5EAED",34)
    connector(d,205,260,22); connector(d,205,350,22); connector(d,690,305,24)
    for y in (250,340): d.line((240,y,515,305),fill="#2E74B5",width=8)
    d.line((515,305,650,305),fill="#2B6E5B",width=8)
    d.text((330,250),"RF",font=fnt(44,True),fill="#566875")
    for x in range(270,620,45): d.line((x,395,x+20,395),fill="#7C8A93",width=3)
    return save_card(im,"03_合路器分路器.png")

def lna_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(240,180,660,430),"#D7DDE1","#929FA7","#F3F5F6",26)
    connector(d,240,305,20); connector(d,660,305,20)
    d.line((280,305,620,305),fill="#2E74B5",width=7)
    d.polygon([(410,340),(450,255),(490,340)],fill="#2B6E5B")
    d.text((360,365),"LNA",font=fnt(42,True),fill="#435866")
    d.line((480,430,540,500),fill="#C23B3B",width=10); d.line((535,500,590,500),fill="#111",width=10)
    return save_card(im,"04_低噪声放大器LNA.png")

def pa_card():
    im,d=shadow_card(); badge(d)
    # Heat-sink enclosure
    box3d(d,(185,195,705,435),"#5C6971","#35414A","#9AA5AB",30)
    for x in range(210,680,25): d.polygon([(x,195),(x+10,145),(x+18,145),(x+18,435),(x+8,435)],fill="#37444C")
    connector(d,185,310,21); connector(d,705,310,21)
    d.text((390,260),"PA",font=fnt(64,True),fill="#EEF2F4")
    d.rounded_rectangle((350,355,545,402),radius=8,fill="#202A30")
    d.ellipse((385,370,405,390),fill="#50C878"); d.ellipse((425,370,445,390),fill="#E0A526")
    d.line((520,435,580,510),fill="#C83434",width=12); d.line((575,510,650,510),fill="#171717",width=12)
    return save_card(im,"05_功率放大器PA.png")

def filter_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(190,190,710,430),"#C6A965","#86713E","#E6D39B",28)
    connector(d,190,310,22); connector(d,710,310,22)
    for x in range(260,660,72):
        d.ellipse((x,238,x+38,276),fill="#765F2E",outline="#493B20",width=3)
        d.line((x+19,276,x+19,380),fill="#8A7136",width=6)
    d.line((235,310,680,310),fill="#EFE1B6",width=6)
    return save_card(im,"06_射频滤波器.png")

def coax_card():
    im,d=shadow_card(); badge(d)
    cx,cy=450,320
    for r in range(175,55,-28):
        d.arc((cx-r,cy-r,cx+r,cy+r),start=15,end=345,fill="#171B1E",width=22)
        d.arc((cx-r,cy-r,cx+r,cy+r),start=15,end=345,fill="#3A4247",width=5)
    d.line((600,210,720,145),fill="#171B1E",width=24); connector(d,735,138,23,"#D4B66E")
    d.line((300,445,180,505),fill="#171B1E",width=24); connector(d,165,512,23,"#D4B66E")
    return save_card(im,"07_射频同轴线缆与接头.png")

def radio_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(150,175,735,445),"#2A3237","#141A1D","#59636A",25)
    d.rounded_rectangle((205,225,495,350),radius=12,fill="#A7D4D2",outline="#111",width=4)
    d.text((230,255),"145.800",font=fnt(45,True),fill="#183940")
    d.ellipse((570,240,690,360),fill="#161B1E",outline="#88949A",width=8)
    d.ellipse((590,260,670,340),outline="#A7B0B5",width=4)
    for i in range(5): d.rounded_rectangle((205+i*60,375,245+i*60,405),radius=5,fill="#555F64")
    d.ellipse((525,385,550,410),fill="#5DC879")
    return save_card(im,"08_双频收发电台.png")

def computer_card():
    im,d=shadow_card(); badge(d)
    # Monitor with generic spectrum UI
    d.rounded_rectangle((155,145,690,440),radius=20,fill="#252D32",outline="#6F7A81",width=5)
    d.rectangle((180,170,665,410),fill="#F7FAFC")
    for y,c in [(230,"#C9DCEB"),(285,"#F0D5A7"),(340,"#CBE5DA")]: d.line((205,y,640,y),fill=c,width=2)
    pts=[]
    for x in range(205,641,8):
        y=315-int(95*abs(sin((x-205)/88))*((x%47)/47))
        pts.append((x,y))
    d.line(pts,fill="#2E74B5",width=4)
    d.line((420,440,420,495),fill="#505A60",width=14); d.rounded_rectangle((320,488,525,512),radius=8,fill="#68747B")
    d.rounded_rectangle((710,255,805,500),radius=12,fill="#353E43",outline="#7B868C",width=4)
    d.ellipse((748,280,768,300),fill="#4CC473")
    return save_card(im,"09_业务控制电脑.png")

def controller_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(150,175,735,445),"#CBD3D8","#8B979F","#EEF2F4",25)
    d.rounded_rectangle((200,220,465,330),radius=12,fill="#162C35",outline="#687B84",width=4)
    d.text((225,240),"AZ  128.6°",font=fnt(27,True),fill="#8BE0D0")
    d.text((225,280),"EL   42.3°",font=fnt(27,True),fill="#8BE0D0")
    for cx,cy,lab in [(575,245,"↑"),(575,345,"↓"),(520,295,"←"),(630,295,"→")]:
        d.ellipse((cx-30,cy-30,cx+30,cy+30),fill="#53646E",outline="#304047",width=3)
        bb=d.textbbox((0,0),lab,font=fnt(28,True)); d.text((cx-(bb[2]-bb[0])/2,cy-(bb[3]-bb[1])/2-4),lab,font=fnt(28,True),fill="white")
    d.ellipse((510,385,535,410),fill="#4DC774"); d.ellipse((555,385,580,410),fill="#E1A72D")
    return save_card(im,"10_旋转器控制器.png")

def control_cable_card():
    im,d=shadow_card(); badge(d)
    colors=["#1A1A1A","#3D5363","#273C2D"]
    for j,(cx,cy) in enumerate([(320,315),(455,300),(585,330)]):
        for r in range(120,35,-24): d.arc((cx-r,cy-r,cx+r,cy+r),20,340,fill=colors[j],width=16)
    # multi-pin connectors
    d.rounded_rectangle((110,420,235,510),radius=18,fill="#4D5960",outline="#20292E",width=4)
    for iy in range(3):
        for ix in range(4): d.ellipse((132+ix*24,440+iy*22,142+ix*24,450+iy*22),fill="#D8C590")
    d.rounded_rectangle((690,150,810,245),radius=18,fill="#535D62",outline="#20292E",width=4)
    for iy in range(3):
        for ix in range(4): d.ellipse((710+ix*24,170+iy*22,720+ix*24,180+iy*22),fill="#D8C590")
    return save_card(im,"11_控制与供电线缆.png")

def surge_card():
    im,d=shadow_card(); badge(d)
    box3d(d,(250,180,650,430),"#BFC8CD","#818D94","#E5E9EB",26)
    connector(d,250,300,23); connector(d,650,300,23)
    d.line((450,430,450,510),fill="#D2B34F",width=18)
    d.polygon([(420,510),(480,510),(450,555)],fill="#C9A33A",outline="#765E1A")
    d.polygon([(430,245),(470,245),(442,310),(478,310),(420,390),(440,325),(405,325)],fill="#E2B329")
    return save_card(im,"12_馈线浪涌保护器.png")

def draw_satellite(d,cx,cy,scale=1):
    d.rectangle((cx-35*scale,cy-25*scale,cx+35*scale,cy+25*scale),fill="#D6A94D",outline="#5D4B25",width=3)
    d.rectangle((cx-130*scale,cy-34*scale,cx-45*scale,cy+34*scale),fill="#4C78A8",outline="#294B6B",width=3)
    d.rectangle((cx+45*scale,cy-34*scale,cx+130*scale,cy+34*scale),fill="#4C78A8",outline="#294B6B",width=3)
    d.line((cx,cy+25*scale,cx+55*scale,cy+85*scale),fill="#555",width=4)
    d.arc((cx+35*scale,cy+60*scale,cx+105*scale,cy+115*scale),0,180,fill="#555",width=4)

def arrow(d, a, b, color="#2E74B5", width=8):
    d.line((a,b),fill=color,width=width)
    x1,y1=a; x2,y2=b; ang=__import__('math').atan2(y2-y1,x2-x1); L=22
    pts=[(x2,y2),(x2-L*cos(ang-.55),y2-L*sin(ang-.55)),(x2-L*cos(ang+.55),y2-L*sin(ang+.55))]
    d.polygon(pts,fill=color)

def diagram_canvas(title, subtitle):
    im=Image.new("RGB",(1800,1000),"#F6F8FA"); d=ImageDraw.Draw(im)
    d.rectangle((0,0,1800,100),fill="#163A5F")
    d.text((70,24),title,font=fnt(44,True),fill="white")
    d.text((70,112),subtitle,font=fnt(25),fill="#667788")
    return im,d

def node(d,box,title,sub="",fill="#FFFFFF",accent="#2E74B5"):
    x1,y1,x2,y2=box; rounded(d,box,fill,"#B9C8D4",4,25)
    d.rectangle((x1,y1,x1+13,y2),fill=accent)
    d.text((x1+35,y1+22),title,font=fnt(30,True),fill="#163A5F")
    if sub: d.multiline_text((x1+35,y1+72),sub,font=fnt(21),fill="#4F6170",spacing=8)

def overall_diagram():
    im,d=diagram_canvas("地面站总体系统架构示意图","空间目标、室外天馈、室内射频、控制软件与保障系统之间的关系")
    draw_satellite(d,900,235,1.15)
    node(d,(90,410,480,650),"室外天馈与跟踪","UV 定向天线\n方位/俯仰旋转器\n馈线及浪涌保护",accent="#2B6E5B")
    node(d,(560,410,950,650),"射频收发链路","合路/分路、滤波\nLNA、PA、双频电台",accent="#936F18")
    node(d,(1030,410,1420,650),"控制与业务处理","业务控制电脑\n轨道预报、多普勒\n任务管理与数据归档",accent="#2E74B5")
    node(d,(510,755,890,920),"旋转器控制","角度指令、驱动\n位置与限位反馈",accent="#647A8A")
    node(d,(1010,755,1390,920),"运行保障","供配电、接地防雷\n网络、授时、备份",accent="#9B5E4A")
    arrow(d,(900,310),(290,405)); arrow(d,(480,530),(555,530)); arrow(d,(950,530),(1025,530))
    arrow(d,(1190,650),(1190,750),"#647A8A"); arrow(d,(890,835),(1005,835),"#647A8A")
    arrow(d,(700,755),(430,650),"#2E74B5"); arrow(d,(1030,610),(930,610),"#2B6E5B")
    p=ASSET/"系统示意图_01_总体架构.png"; im.save(p); return p

def rf_diagram():
    im,d=diagram_canvas("UV 地面站射频收发链路示意图","接收链路以低噪声和抗干扰为重点；发射链路以功率控制、滤波和联锁为重点")
    d.text((70,185),"接收链路",font=fnt(34,True),fill="#2B6E5B")
    xs=[70,345,620,895,1170,1445]
    names=[("UV 天线","空间下行"),("合路/分路器","频段汇合"),("接收滤波器","抑制带外干扰"),("LNA","低噪声增益"),("双频电台/接收机","调谐与解调"),("控制电脑","频谱、解码、记录")]
    for x,(t,s) in zip(xs,names): node(d,(x,235,x+235,405),t,s,accent="#2B6E5B")
    for i in range(5): arrow(d,(xs[i]+235,320),(xs[i+1]-5,320),"#2B6E5B",7)
    d.text((70,520),"发射链路（受控）",font=fnt(34,True),fill="#9B1C1C")
    names2=[("控制电脑","权限与多普勒"),("双频电台","调制与激励"),("PA","受控功率放大"),("发射滤波器","谐波/带外抑制"),("合路/分路器","馈入目标频段"),("UV 天线","空间上行")]
    for x,(t,s) in zip(xs,names2): node(d,(x,580,x+235,750),t,s,accent="#9B1C1C")
    for i in range(5): arrow(d,(xs[i]+235,665),(xs[i+1]-5,665),"#9B1C1C",7)
    rounded(d,(350,840,1450,940),"#FFF2F1","#C66A65",3,18)
    d.text((390,867),"发射默认禁用：资质、许可、频率/模式、功率、角度与设备状态均满足后方可开放。",font=fnt(26,True),fill="#8B2727")
    p=ASSET/"系统示意图_02_射频链路.png"; im.save(p); return p

def control_diagram():
    im,d=diagram_canvas("自动跟踪与设备控制关系示意图","业务控制电脑同时协调天线指向、频率补偿、录制解码与任务日志")
    node(d,(80,205,420,380),"任务与轨道数据","目标卫星、站点坐标\n轨道根数、任务模板",accent="#647A8A")
    node(d,(600,205,1100,440),"业务控制电脑","过境预报\n方位/俯仰计算\n多普勒补偿\n任务编排与数据归档",accent="#2E74B5")
    node(d,(1330,175,1720,360),"双频电台/接收机","调谐、收发控制\n音频/基带/频谱数据",accent="#936F18")
    node(d,(120,650,500,865),"旋转器控制器","目标角度接收\n电机驱动\n限位与故障状态",accent="#647A8A")
    node(d,(710,650,1090,865),"方位/俯仰旋转器","二维运动\n角度位置反馈",accent="#2B6E5B")
    node(d,(1300,650,1680,865),"UV 双频天线","跟随卫星过境\n完成空间收发",accent="#2B6E5B")
    arrow(d,(420,292),(595,292),"#647A8A")
    arrow(d,(1100,290),(1325,290),"#936F18")
    arrow(d,(750,440),(350,645),"#2E74B5")
    arrow(d,(500,755),(705,755),"#2E74B5")
    arrow(d,(1090,755),(1295,755),"#2B6E5B")
    arrow(d,(770,650),(470,570),"#647A8A")
    d.text((495,535),"实际角度 / 限位 / 故障反馈",font=fnt(24,True),fill="#647A8A")
    arrow(d,(1450,650),(1450,365),"#936F18")
    d.text((1475,480),"射频信号",font=fnt(24,True),fill="#936F18")
    p=ASSET/"系统示意图_03_跟踪控制.png"; im.save(p); return p

def set_run(run,size=10.5,bold=False,color=DARK,italic=False):
    run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)

def set_alt(shape,text):
    shape._inline.docPr.set("descr",text); shape._inline.docPr.set("title",text)

def cell_margins(cell,top=100,start=120,bottom=100,end=120):
    tcPr=cell._tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in("w:tcMar")
    if mar is None: mar=OxmlElement("w:tcMar"); tcPr.append(mar)
    for tag,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        e=mar.find(qn(f"w:{tag}"))
        if e is None: e=OxmlElement(f"w:{tag}"); mar.append(e)
        e.set(qn("w:w"),str(val)); e.set(qn("w:type"),"dxa")

def borders(table,color="CBD6DF",size="6"):
    pr=table._tbl.tblPr; bd=pr.find(qn("w:tblBorders"))
    if bd is None: bd=OxmlElement("w:tblBorders"); pr.append(bd)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e=OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),size); e.set(qn("w:color"),color); bd.append(e)

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); r=p.add_run(text); set_run(r,{1:16,2:13}[level],True,BLUE); return p

def add_caption(doc,text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(8)
    set_run(p.add_run(text),8.5,False,MUTED,True)

MODULES = [
    ("UV 双频定向天线", "现场实景", "分别承担 VHF/UHF 频段的方向性收发；单元数、极化和安装间距由链路预算确定。", "01_UV双频定向天线_现场实景.png"),
    ("方位/俯仰旋转器", "现场实景", "带动天线完成方位和俯仰二维运动，具备位置反馈、限位和安全归位能力。", "02_方位俯仰旋转器_现场实景.png"),
    ("合路器/分路器", "典型外观示意", "按频段合并或分离射频通道，接口数量、隔离度、功率能力需按系统方案选定。", "03_合路器分路器.png"),
    ("低噪声放大器（LNA）", "典型外观示意", "靠近接收前端提升弱信号，重点关注噪声系数、增益、过载和偏置供电。", "04_低噪声放大器LNA.png"),
    ("功率放大器（PA）", "典型外观示意", "对发射激励进行受控放大，需配置功率限制、散热、滤波和发射联锁。", "05_功率放大器PA.png"),
    ("射频滤波器", "典型外观示意", "对工作频段进行带通或陷波处理，抑制带外干扰、谐波和共址信号影响。", "06_射频滤波器.png"),
    ("射频同轴线缆与接头", "典型外观示意", "连接天线与射频设备；线径、长度、损耗、屏蔽、防水和弯曲半径需统一核算。", "07_射频同轴线缆与接头.png"),
    ("双频收发电台", "典型外观示意", "完成频率调谐、语音/数据调制解调和受控收发，并接受多普勒修正。", "08_双频收发电台.png"),
    ("业务控制电脑", "典型外观示意", "运行轨道预报、自动跟踪、多普勒、频谱/解码、任务管理和数据归档软件。", "09_业务控制电脑.png"),
    ("旋转器控制器", "典型外观示意", "接收方位/俯仰目标角，驱动旋转器并返回实际角度、限位和故障状态。", "10_旋转器控制器.png"),
    ("控制与供电线缆", "典型外观示意", "传输电机驱动、角度反馈、设备控制和供电信号；须编号、固定并预留运动余量。", "11_控制与供电线缆.png"),
    ("馈线浪涌保护器", "典型外观示意", "在馈线入室或机箱入口处泄放浪涌并接入等电位系统，不能替代完整防雷设计。", "12_馈线浪涌保护器.png"),
]

def add_gallery(doc, items, page_title):
    heading=add_heading(doc,page_title,1)
    heading.paragraph_format.page_break_before=True
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(7)
    set_run(p.add_run("说明：除明确标注“现场实景”的项目外，其余均为无品牌的典型实物形态插图，不代表具体型号、尺寸、接口数量或最终供货形态。"),9.2,False,MUTED,True)
    table=doc.add_table(rows=3,cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; borders(table)
    for idx,(title,kind,desc,filename) in enumerate(items):
        cell=table.cell(idx//2,idx%2); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP; cell.width=Inches(3.22); cell_margins(cell)
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
        shape=p.add_run().add_picture(str(ASSET/filename),width=Inches(2.52)); set_alt(shape,f"{title}：{kind}")
        hp=cell.add_paragraph(); hp.paragraph_format.space_before=Pt(1); hp.paragraph_format.space_after=Pt(2); hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_run(hp.add_run(title),10.2,True,NAVY)
        dp=cell.add_paragraph(); dp.paragraph_format.space_after=Pt(0); dp.paragraph_format.line_spacing=1.05
        set_run(dp.add_run(desc),8.2,False,DARK)

def build():
    photo_card(PHOTO1,(0.14,0.19,0.86,0.67),"01_UV双频定向天线_现场实景.png")
    photo_card(PHOTO2,(0.37,0.22,0.72,0.70),"02_方位俯仰旋转器_现场实景.png")
    combiner_card(); lna_card(); pa_card(); filter_card(); coax_card(); radio_card(); computer_card(); controller_card(); control_cable_card(); surge_card()
    diagrams=[overall_diagram(),rf_diagram(),control_diagram()]

    doc=Document(SRC)
    # Move the existing end marker to the new end.
    for p in list(doc.paragraphs):
        if "文档结束" in p.text:
            p._element.getparent().remove(p._element)

    doc.add_page_break(); add_heading(doc,"附录 C  系统示意图与模块实物图册",1)
    p=doc.add_paragraph(); set_run(p.add_run("本附录用于帮助甲方直观理解系统组成和设备外观。系统示意图表达功能连接关系，不作为施工接线图；模块图片用于说明典型实物形态，最终接口、性能、尺寸和安装方式应以深化设计及供货资料为准。"),10.5)
    add_heading(doc,"C.1 地面站总体系统架构",2)
    shape=doc.add_picture(str(diagrams[0]),width=Inches(6.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; set_alt(shape,"地面站总体系统架构示意图")
    add_caption(doc,"图 C-1  地面站总体系统架构示意图")
    add_heading(doc,"C.2 射频接收与发射链路",2)
    shape=doc.add_picture(str(diagrams[1]),width=Inches(6.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; set_alt(shape,"UV地面站射频接收与受控发射链路示意图")
    add_caption(doc,"图 C-2  射频接收与受控发射链路示意图")

    doc.add_page_break(); add_heading(doc,"C.3 自动跟踪与设备控制关系",2)
    shape=doc.add_picture(str(diagrams[2]),width=Inches(6.4)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; set_alt(shape,"自动跟踪与设备控制关系示意图")
    add_caption(doc,"图 C-3  自动跟踪与设备控制关系示意图")
    add_heading(doc,"C.4 模块连接关系说明",2)
    rows=[
        ("接收主链路","UV 天线 → 合路/分路器 → 接收滤波器 → LNA → 双频电台/接收机 → 业务控制电脑"),
        ("发射主链路","业务控制电脑（授权/多普勒）→ 双频电台 → PA → 发射滤波器 → 合路/分路器 → UV 天线"),
        ("跟踪控制链路","业务控制电脑 → 旋转器控制器 → 方位/俯仰旋转器 → UV 天线；位置、限位和故障状态反向反馈"),
        ("保障链路","供配电、控制/供电线缆、馈线浪涌保护、接地与网络授时共同保障系统运行"),
    ]
    table=doc.add_table(rows=1,cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; borders(table)
    for i,h in enumerate(("链路","典型连接关系")):
        cell=table.rows[0].cells[i]; cell_margins(cell); cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd")); cell._tc.tcPr[-1].set(qn("w:fill"),"EAF2F8")
        hp=cell.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run(hp.add_run(h),9.5,True,NAVY)
    for a,b in rows:
        cells=table.add_row().cells
        for i,val in enumerate((a,b)):
            cell_margins(cells[i]); set_run(cells[i].paragraphs[0].add_run(val),9.0, i==0, DARK)

    add_gallery(doc,MODULES[:6],"附录 D  室外与射频模块典型实物形态")
    add_gallery(doc,MODULES[6:],"附录 E  收发、控制与线缆模块典型实物形态")

    doc.core_properties.title="业余无线电卫星UV通信地面站建设方案及功能介绍（图示增强版）"
    doc.save(OUT)
    print(OUT)

if __name__=="__main__": build()
